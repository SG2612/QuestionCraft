from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pymysql
from pymysql.err import MySQLError
import tkinter as tk
from tkinter import filedialog, simpledialog

# Function to create DOCX
def create_docx(data, path):
    if not data:
        print("No data to generate DOCX.")
        return

    class_info, mcq_questions, saq_questions, long_questions, broad_questions, mcq_answers = data

    # Calculate marks for each question type
    saq_marks = len(saq_questions) * 2
    mcq_marks = len(mcq_questions)
    long_marks = len(long_questions) * 5
    broad_marks = len(broad_questions) * 10

    # Calculate total marks
    total_marks = saq_marks + mcq_marks + long_marks + broad_marks

    # Create a new Document
    doc = Document()

    # Add the title
    title = doc.add_heading('XYZ School', level=1)
    title.alignment = 1  # Center align the title

    # Add class information
    class_paragraph = doc.add_paragraph(f'Class: {class_info}')
    class_paragraph.alignment = 1  # Center align

    # Add total marks
    marks_paragraph = doc.add_paragraph(f'Full marks: {total_marks}')
    marks_paragraph.alignment = 1  # Center align

    # Add a horizontal line
    doc.add_paragraph('')

    # Function to add questions
    def add_questions(heading, questions, marks, question_type=None):
        doc.add_heading(f'Answer the following {heading} ({marks} marks):', level=2)
        for idx, question in enumerate(questions, start=1):
            question_paragraph = doc.add_paragraph(f'{idx}. {question[0]}')
            if question_type == 'mcq':
                for option_label, option in zip(['A', 'B', 'C', 'D'], question[1:]):
                    option_paragraph = doc.add_paragraph(f'{option_label}. {option}')
                    option_paragraph.paragraph_format.left_indent = Pt(36)

    # Add the questions with separate numbering for each type
    add_questions('MCQ questions', mcq_questions, mcq_marks, question_type='mcq')
    doc.add_paragraph('')  # Add space before the next section
    add_questions('SAQ questions', saq_questions, saq_marks)
    doc.add_paragraph('')  # Add space before the next section
    add_questions('Long questions', long_questions, long_marks)
    doc.add_paragraph('')  # Add space before the next section
    add_questions('Broad questions', broad_questions, broad_marks)

    # Add a new page for MCQ answers
    doc.add_page_break()
    doc.add_heading('MCQ Answers', level=2)
    for idx, answer in enumerate(mcq_answers, start=1):
        doc.add_paragraph(f'{idx}. {answer}')

    # Save the DOCX
    doc.save(path)

# Fetch data from MySQL database
def fetch_data():
    data = []
    my_conn = None

    try:
        # MySQL connection
        my_conn = pymysql.connect(
            host='localhost',
            user='root',
            password='krishnendu@2003',
            database='project'
        )

        with my_conn.cursor() as cursor:
            # Fetch class information
            cursor.execute("SELECT class FROM final LIMIT 1")
            class_info = cursor.fetchone()[0]

            # Fetch questions and answers
            question_types = ['mcq', 'saq', 'long', 'broad']
            questions = {q_type: [] for q_type in question_types}
            mcq_answers = []

            for q_type in question_types:
                if q_type == 'mcq':
                    cursor.execute(f"SELECT QUESTION, A, B, C, D, ANSWER FROM final WHERE QUESTION_TYPE='{q_type}'")
                    questions[q_type] = [row[:5] for row in cursor.fetchall()]
                    cursor.execute(f"SELECT ANSWER FROM final WHERE QUESTION_TYPE='{q_type}'")
                    mcq_answers = [row[0] for row in cursor.fetchall()]
                else:
                    cursor.execute(f"SELECT QUESTION FROM final WHERE QUESTION_TYPE='{q_type}'")
                    questions[q_type] = [(row[0],) for row in cursor.fetchall()]

            data = (class_info, questions['mcq'], questions['saq'], questions['long'], questions['broad'], mcq_answers)
            print("Data fetched from database:", data)

    except MySQLError as e:
        print("Error: ", e)

    finally:
        if my_conn:
            my_conn.close()

    return data

# Function to select folder and get file name
def select_folder_and_filename():
    root = tk.Tk()
    root.withdraw()  # Hide the root window

    folder_selected = filedialog.askdirectory()
    if folder_selected:
        file_name = simpledialog.askstring("Input", "Enter file name:", parent=root)
        if file_name:
            return folder_selected, file_name
    return None, None

# Main execution
if __name__ == "__main__":
    data = fetch_data()
    if data:
        folder_path, file_name = select_folder_and_filename()
        if folder_path and file_name:
            docx_path = f"{folder_path}/{file_name}.docx"
            create_docx(data, docx_path)
        else:
            print("No folder or file name selected.")
    else:
        print("No data fetched to create the document.")
